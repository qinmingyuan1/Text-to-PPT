from docx import Document
import markdown
from markdown.treeprocessors import Treeprocessor
from markdown.extensions import Extension
from collections import defaultdict
from bs4 import BeautifulSoup
import pandas as pd

class UnsupportedFileTypeError(Exception):
    def __init__(self, file_type):
        self.file_type = file_type
        self.message = f"Unsupported file type: {file_type}"
        super().__init__(self.message)


def parse(my_file):
    content = None
    if my_file.type == 'text/csv':
        content = {}
        content['type'] = 'csv'
        content['text'] = pd.read_csv(my_file, index_col=False).to_string(index=False)
    elif my_file.type == 'application/octet-stream':
        content = parse_markdown(my_file)
    elif my_file.type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        content = parse_docx(my_file)
    elif my_file.type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
        content = parse_xlsx(my_file)
    elif my_file.type == 'text/plain':
        content = {}
        content['type'] = 'txt'
        content['text'] = my_file.getvalue().decode("utf-8")
    else:
        raise UnsupportedFileTypeError(my_file.type)

    return content

def parse_docx(my_file):
    content = {
        'type': 'docx',
        'text': '',
        'resources': {'images': [], 'charts': [], 'tables': []}
    }
    doc = Document(my_file)
    
    paragraph_counter = 1
    for para in doc.paragraphs:
        print(para.style.name)
        if para.style.name.startswith('Title'):
            content['text'] += f'tile: {para.text}\n'
        elif para.style.name.startswith('Subtitle'):
            content['text'] += f'subtile: {para.text}\n'
        elif para.style.name.startswith('Heading'):
            level = int(para.style.name.split()[-1])
            heading_key = f'heading_{level}'
            content['text'] += f'{heading_key}: {para.text}\n'
            current_key = heading_key
        else:
            paragraph_key = f'paragraphs_{paragraph_counter}'
            content['text'] += f'{paragraph_key}: {para.text}\n'
            current_key = paragraph_key
            paragraph_counter += 1

    # 解析图片和图表
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            content['resources']['images'].append(rel.target_ref)
        elif "chart" in rel.target_ref:
            content['resources']['charts'].append(rel.target_ref)

    # 解析表格
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        content['resources']['tables'].append(table_data)

    return content


class MarkdownProcessor(Treeprocessor):
    def run(self, root):
        # 直接返回 root，不进行处理
        return root

class MarkdownExtension(Extension):
    def extendMarkdown(self, md):
        md.treeprocessors.register(MarkdownProcessor(md), 'markdown_processor', 15)

def parse_markdown(my_file):
    content = {
        'type': 'md',
        'text': '',
        'resources': {'images': [], 'charts': [], 'tables': []}
    }
    md = markdown.Markdown(extensions=[MarkdownExtension()])
    html = md.convert(my_file.read().decode('utf-8'))
    soup = BeautifulSoup(html, 'html.parser')

    paragraph_counter = 1
    for element in soup.descendants:
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(element.name[1])
            heading_key = f'heading_{level}'
            content['text'] += f'{heading_key}: {element.text}\n'
        elif element.name == 'p':
            paragraph_key = f'paragraphs_{paragraph_counter}'
            content['text'] += f'{paragraph_key}: {element.text}\n'
            paragraph_counter += 1
        elif element.name == 'img':
            content['resources']['images'].append(element['src'])
        elif element.name == 'table':
            table_data = []
            for row in element.find_all('tr'):
                row_data = [cell.text for cell in row.find_all(['th', 'td'])]
                table_data.append(row_data)
            content['resources']['tables'].append(table_data)

    return content


def parse_xlsx(my_file):
    content = {
        'type': 'xlsx',
        'text': '',
        'resources': {'images': [], 'charts': [], 'tables': []}
    }
    xls = pd.ExcelFile(my_file)
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name)
        content['text'] += f'Sheet: {sheet_name}\n{df.to_string(index=False)}\n'
        if sheet_name not in content['resources']:
            content['resources'][f'Sheet: {sheet_name}'] = {'tables': []}
        content['resources'][f'Sheet: {sheet_name}']['tables'].append(df.values.tolist())
    return content
